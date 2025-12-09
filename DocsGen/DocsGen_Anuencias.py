import os
import tkinter as tk
import comtypes.client
import re
import traceback
import uuid
import time
import glob
import json
from plyer import notification
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from ttkbootstrap import Style
from PIL import Image, ImageTk
from datetime import datetime
import threading
import pythoncom

# Obter a data de hoje
hoje = datetime.today()
# Formatar a data como DDMMAAAA
data_hoje = hoje.strftime("%d%m%Y")

# Obtém o diretório onde o script está localizado
caminho_base = os.path.dirname(os.path.abspath(__file__))

# Lista de meses em português
meses = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

# Lista de documentos disponíveis (Unificada)
DOCS_DISPONIVEIS = ["NR10", "NR10 SEP", "NR12", "NR33", "NR35"]

class genAnuencias:
    def __init__(self, root):
        # Crie a interface gráfica
        style = Style(theme='darkly')
        self.root = root
        root = style.master
        root.title("Gerador de Anuências - Unificado")
        frame = ttk.Frame(root)
        frame['padding'] = (10, 10, 10, 10)
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        self.dir_pasta_selecionada = ""

        # Carregar a imagem
        try:
            self.bg_image = Image.open(r"C:\PyProjects\DocsGen\bkg_anuencias.png")
            self.bg_photo = ImageTk.PhotoImage(self.bg_image)
            ttk.Label(frame, image=self.bg_photo).grid(row=0, column=0, columnspan=6, pady=(0, 20))
        except Exception:
            pass

        #LabelFrame - Dados Documento
        lblframe_dados = ttk.LabelFrame(frame, text="Dados Documento:", padding=10)
        lblframe_dados.grid(row=1, column=0, columnspan=6, sticky="ew", padx=10, pady=5)   

        ttk.Label(lblframe_dados, text="Funcionário: ").grid(row=0, column=0, sticky=tk.W)
        self.entr_funcionario = ttk.Entry(lblframe_dados, width=60)
        self.entr_funcionario.grid(row=0, column=1, columnspan=4, sticky=tk.W, pady=5)

        ttk.Label(lblframe_dados, text="CPF: ").grid(row=1, column=0, sticky=tk.W)
        self.entr_CPF = ttk.Entry(lblframe_dados, width=30)
        self.entr_CPF.grid(row=1, column=1, sticky=tk.W, pady=5)

        ttk.Label(lblframe_dados, text="Iniciais: ").grid(row=1, column=2, sticky=tk.W)
        self.entr_apelido = ttk.Entry(lblframe_dados, width=15)
        self.entr_apelido.grid(row=1, column=3, sticky=tk.W, pady=5)

        ttk.Label(lblframe_dados, text="HSE Responsável: ").grid(row=2, column=0, sticky=tk.W)
        self.cbbx_hse = ttk.Combobox(lblframe_dados, values=[], width=30)
        self.cbbx_hse.grid(row=2, column=1, sticky=tk.W, pady=5)

        # Seleção de Cargo
        lblframe_cargo = ttk.LabelFrame(frame, text="Seleção de Cargo:", padding=10)
        lblframe_cargo.grid(row=2, column=0, columnspan=6, sticky="ew", padx=10, pady=5)
        
        ttk.Label(lblframe_cargo, text="GHE: ").grid(row=0, column=0, sticky=tk.W)
        self.cbbx_ghe = ttk.Combobox(lblframe_cargo, values=["01", "02", "03", "04", "05"], width=10, state="readonly")
        self.cbbx_ghe.grid(row=0, column=1, sticky=tk.W, padx=10)
        self.cbbx_ghe.bind("<<ComboboxSelected>>", self.atualizar_cargos_por_ghe)

        ttk.Label(lblframe_cargo, text="Cargo: ").grid(row=0, column=2, sticky=tk.W)
        self.combo_cargo = ttk.Combobox(lblframe_cargo, values=[], state="readonly", width=40)
        self.combo_cargo.grid(row=0, column=3, sticky=tk.W, padx=10)
        self.combo_cargo.bind("<<ComboboxSelected>>", self.atualizar_opcoes_cargo)

        # Frame para Checkboxes (Dinâmico)
        self.lblframe_anuencias = ttk.LabelFrame(frame, text="Anuências Disponíveis:", padding=10)
        self.lblframe_anuencias.grid(row=3, column=0, columnspan=6, sticky="ew", padx=10, pady=5)
        
        self.vars_docs = {} # Dicionário para guardar as variáveis dos checkboxes

        # Opções
        lblframe_opcoes = ttk.LabelFrame(frame, text="Opções: ", padding=10)
        lblframe_opcoes.grid(row=4, column=0, columnspan=6, sticky="ew", padx=10, pady=5)                

        ttk.Label(lblframe_opcoes, text="Salvar em: ").grid(row=0, column=0, sticky=tk.W)      
        ttk.Button(lblframe_opcoes, text="...", command=self.selecionar_pasta).grid(row=0, column=1, sticky=tk.W, padx=(0, 10))   
        self.lbl_pastaselecionada = tk.Label(lblframe_opcoes, text="Selecione a pasta", wraplength=350)
        self.lbl_pastaselecionada.grid(row=0, column=2, columnspan=4, sticky=tk.W, padx=(0,10))

        # Botão Gerar
        ttk.Button(frame, text="Gerar Anuências", command=self.verificar_checkbuttons).grid(row=5, column=0, sticky=tk.W, padx=10, pady=20)
        self.lbl_avisoGeracao = ttk.Label(frame, text="")
        self.lbl_avisoGeracao.grid(row=5, column=1, columnspan=4, sticky=tk.W)

        self.dados = {}

        # Carregar HSEs
        self.lista_hses = self.carregar_hses()
        self.cbbx_hse['values'] = list(self.lista_hses.keys())

        # Carregar GHEs do JSON
        self.ghe_data = self.carregar_ghe_json()

    def carregar_ghe_json(self):
        arquivo = os.path.join(caminho_base, 'ghe_config.json')
        if os.path.exists(arquivo):
            try:
                with open(arquivo, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"Erro ao ler ghe_config.json: {e}")
        return {}

    def carregar_hses(self):
        hses = {}
        arquivo = os.path.join(caminho_base, 'listaHSE.txt')
        if os.path.exists(arquivo):
            try:
                with open(arquivo, 'r', encoding='utf-8') as file:
                    for line in file:
                        parts = line.strip().split(';')
                        if len(parts) >= 3:
                            # ROLE;NAME;REGISTRATION
                            role = parts[0].strip()
                            name = parts[1].strip()
                            reg = parts[2].strip()
                            
                            # Mapeamento de função se necessário
                            funcao = "Técnico(a) de Segurança do Trabalho" if role.upper() == "HSE" else role
                            
                            hses[name] = {
                                "nome": name,
                                "registro": reg,
                                "funcao": funcao
                            }
            except Exception as e:
                print(f"Erro ao ler listaHSE.txt: {e}")
        return hses
        
    def atualizar_cargos_por_ghe(self, event):
        gheSelecionado = self.cbbx_ghe.get()
        if gheSelecionado in self.ghe_data:
            self.combo_cargo['values'] = self.ghe_data[gheSelecionado]
        else:
            self.combo_cargo['values'] = []
        
        self.combo_cargo.set('') # Limpar seleção anterior
        # Limpar checkboxes de anuencias
        for widget in self.lblframe_anuencias.winfo_children():
            widget.destroy()
        self.vars_docs.clear()

    def carregar_cargos_do_arquivo(self):
        cargos = []
        arquivo = os.path.join(caminho_base, 'listaDescFuncoes.txt')
        if os.path.exists(arquivo):
            try:
                with open(arquivo, 'r', encoding='utf-8') as file:
                    for line in file:
                        parts = line.strip().split(';')
                        if len(parts) > 1:
                            cargos.append(parts[1].strip())
            except Exception as e:
                print(f"Erro ao ler arquivo de funções: {e}")
        return sorted(list(set(cargos))) # Remove duplicatas e ordena

    def atualizar_opcoes_cargo(self, event):
        # Limpar checkboxes antigos
        for widget in self.lblframe_anuencias.winfo_children():
            widget.destroy()
        self.vars_docs.clear()
        
        # Exibir todas as opções disponíveis
        for i, doc in enumerate(DOCS_DISPONIVEIS):
            var = tk.BooleanVar()
            chk = ttk.Checkbutton(self.lblframe_anuencias, text=doc, variable=var)
            chk.grid(row=0, column=i, padx=10, sticky=tk.W)
            self.vars_docs[doc] = var

    def selecionar_pasta(self):
        pasta_selecionada = filedialog.askdirectory()
        if pasta_selecionada:
            self.lbl_pastaselecionada.config(text=pasta_selecionada)   
            self.dir_pasta_selecionada = pasta_selecionada  

    def substituir_texto(self, paragraph, substituicoes):
        # Substituir o texto nos runs do parágrafo
        for palavra_antiga, palavra_nova in substituicoes.items():
            if palavra_nova is None:
                palavra_nova = ""
            if palavra_antiga in paragraph.text:
                if palavra_antiga == 'NOMEFUNCIONARIO':
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                            run.bold = True  # Mantém o negrito
                        run.font.name = 'Arial'
                else:
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                        run.font.name = 'Arial'
        return paragraph

    def substituir_texto_tabela(self, cell, substituicoes):
        # Substituir o texto nas células da tabela
        for paragraph in cell.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)  # Chamando a função que substitui no parágrafo
        # Centralizar verticalmente o conteúdo da célula
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return cell

    def validar_cpf(self, cpf: str) -> bool:
            """Valida um CPF informado com pontos e traço."""

            # Remover caracteres não numéricos (pontos e traço)
            cpf = re.sub(r"[\.\-]", "", cpf)

            # Verificar se tem exatamente 11 dígitos
            if not cpf.isdigit() or len(cpf) != 11:
                return False

            # Verificar se todos os dígitos são iguais (ex: 111.111.111-11 é inválido)
            if cpf == cpf[0] * 11:
                return False

            # Cálculo do primeiro dígito verificador
            soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
            digito1 = (soma * 10 % 11) % 10

            # Cálculo do segundo dígito verificador
            soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
            digito2 = (soma * 10 % 11) % 10

            # Verifica se os dígitos calculados são iguais aos do CPF informado
            return digito1 == int(cpf[9]) and digito2 == int(cpf[10])

    def mostrar_progresso(self):
        """Exibe barra de progresso e inicia a geração em thread"""
        self.janela_progress = tk.Toplevel(self.root)
        self.janela_progress.title("Gerando Documentos...")
        self.janela_progress.geometry("400x150")
        self.janela_progress.resizable(False, False)
        self.janela_progress.transient(self.root)
        self.janela_progress.grab_set()
        
        # Centralizar
        self.janela_progress.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (self.janela_progress.winfo_width() // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (self.janela_progress.winfo_height() // 2)
        self.janela_progress.geometry(f"+{x}+{y}")
        
        ttk.Label(self.janela_progress, text="Gerando documentos...\nPor favor, aguarde...", justify=tk.CENTER).pack(pady=20)
        
        self.progress = ttk.Progressbar(self.janela_progress, mode='indeterminate', length=350)
        self.progress.pack(pady=10, padx=20)
        self.progress.start()
        
        # Iniciar thread
        threading.Thread(target=self._executar_geracao, daemon=True).start()

    def limpar_temporarios(self, pasta):
        """Remove arquivos temporários da pasta especificada."""
        try:
            padrao = os.path.join(caminho_base, pasta, "temp_*.docx")
            arquivos = glob.glob(padrao)
            for arquivo in arquivos:
                # Tentar deletar com retentativa
                for _ in range(3):
                    try:
                        os.remove(arquivo)
                        break
                    except Exception:
                        time.sleep(0.5)
        except Exception:
            pass

    def _executar_geracao(self):
        pasta_templates = "templates"
        try:
            pythoncom.CoInitialize()
            cargo_selecionado = self.combo_cargo.get().upper()
            
            # Gerar sufixo baseado no nome do cargo (sanitizado)
            # Ex: "Técnico O&M" -> "_TECNICO_O_M"
            sufixo = "_" + re.sub(r'[^a-zA-Z0-9]', '_', cargo_selecionado).upper()
            
            # Limpar temporários antigos antes de começar
            self.limpar_temporarios(pasta_templates)
            
            # Preparar dados para substituição
            nome_func = self.entr_funcionario.get().upper()
            cpf_func = self.entr_CPF.get().upper()
            
            # Garantir que não sejam None
            if nome_func is None: nome_func = ""
            if cpf_func is None: cpf_func = ""

            # Lógica HSE
            hse_selecionado = self.cbbx_hse.get().upper()
            nomeHSE = ""
            registroHSE = ""
            funcaoHSE = ""

            if hse_selecionado in self.lista_hses:
                dados_hse = self.lista_hses[hse_selecionado]
                nomeHSE = dados_hse["nome"]
                registroHSE = dados_hse["registro"]
                funcaoHSE = dados_hse["funcao"]
            elif hse_selecionado: # Caso tenha digitado manualmente ou algo assim (fallback simples)
                 nomeHSE = hse_selecionado

            self.dados = {
                'NOMEFUNCIONARIO': str(nome_func),
                'FUNCAOFUNCIONARIO': str(cargo_selecionado),
                'DIAANUENCIA': datetime.today().strftime("%d"),
                'MESANUENCIA': str(meses.get(datetime.today().month, datetime.today().strftime("%B"))),
                'ANOANUENCIA': datetime.today().strftime("%Y"),
                'CPFFUNCIONARIO': str(cpf_func),
                'NOMEHSE': nomeHSE,
                'REGISTROHSE': registroHSE,
                'FUNCAOHSE': funcaoHSE
            }

            for doc_nome, var in self.vars_docs.items():
                if var.get():
                    # Determinar nome do template
                    # Ex: NR10 -> NR10.docx
                    # Ex: NR10 SEP -> NR10_SEP.docx
                    nome_arquivo_template = doc_nome.replace(" ", "_") + ".docx"
                    
                    self.gerar_documento(nome_arquivo_template, pasta_templates, sufixo, doc_nome)
                    
                    self.root.after(0, lambda d=doc_nome: notification.notify(
                        title="Aviso",
                        message=f"{d} Gerada com Sucesso!",
                        timeout=5
                    ))

            # Finalização
            self.root.after(0, lambda: [
                self.janela_progress.destroy(),
                self.lbl_avisoGeracao.config(text=""),
                messagebox.showinfo("Concluído", "Documento(s) Salvo(s) com Sucesso!")
            ])

        except Exception as e:
            erro_msg = traceback.format_exc()
            self.root.after(0, lambda: [
                self.janela_progress.destroy(),
                messagebox.showerror("Erro", f"Ocorreu um erro:\n{erro_msg}")
            ])
        finally:
            self.limpar_temporarios(pasta_templates)
            pythoncom.CoUninitialize()

    def gerar_documento(self, template_name, pasta, sufixo, doc_nome):
        template_path = os.path.join(caminho_base, pasta, template_name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        doc = Document(template_path)

        for paragraph in doc.paragraphs:
            self.substituir_texto(paragraph, self.dados)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.substituir_texto_tabela(cell, self.dados)

        # Nome do arquivo de saída
        # Ex: NR10_Apelido_Data_Sufixo.pdf
        prefixo_saida = doc_nome.replace(" ", "_") # NR10, NR10_SEP
        
        # Validar caracteres inválidos no apelido
        apelido = self.entr_apelido.get().upper()
        if apelido is None: apelido = ""
        apelido = re.sub(r'[<>:"/\\|?*]', '', apelido) # Removes caracteres inválidos
        
        nome_arquivo = f"{prefixo_saida}_{apelido}_{data_hoje}{sufixo}.pdf"
        
        # Gerar nome temporário único para evitar conflitos de arquivo em uso
        unique_id = uuid.uuid4().hex
        temp_docx = os.path.join(caminho_base, pasta, f"temp_{prefixo_saida}_{unique_id}.docx")
        pdf_path = os.path.join(self.dir_pasta_selecionada, nome_arquivo)

        # Verificar se o PDF de destino já está aberto (tentativa simples de abrir para escrita)
        if os.path.exists(pdf_path):
            try:
                with open(pdf_path, 'ab'):
                    pass
            except PermissionError:
                raise PermissionError(f"O arquivo PDF '{nome_arquivo}' parece estar aberto em outro programa. Feche-o e tente novamente.")

        doc.save(temp_docx)

        word = None
        doc_word = None
        try:
            word = comtypes.client.CreateObject('Word.Application')
            # word.Visible = False # Opcional, mas bom para evitar janelas piscando
            doc_word = word.Documents.Open(os.path.abspath(temp_docx))
            
            # Tentar salvar, capturando erros específicos do Word/COM
            try:
                doc_word.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            except Exception as e:
                if "permission" in str(e).lower() or "access is denied" in str(e).lower():
                     raise PermissionError(f"Erro de permissão ao salvar '{nome_arquivo}'. Verifique se o arquivo está aberto.")
                raise e
                
        except Exception as e:
            raise e
        finally:
            if doc_word:
                try:
                    doc_word.Close()
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass
            
            # Garantir remoção do arquivo temporário
            if os.path.exists(temp_docx):
                try:
                    time.sleep(0.5) # Pequena pausa para liberar o arquivo
                    os.remove(temp_docx)
                except:
                    pass
            
    def verificar_checkbuttons(self):      

        cpf_digitado = self.entr_CPF.get()
        texto_pastaselecionada = self.lbl_pastaselecionada.cget("text")
        cargo = self.combo_cargo.get()

        if not cargo:
            messagebox.showinfo("Alerta!", "Selecione um Cargo!")
            return

        if texto_pastaselecionada == "Selecione a pasta":
            messagebox.showinfo("Alerta!", "Selecione a pasta para salvar o(s) documento(s)!") 
        else:
            if self.validar_cpf(cpf_digitado):    
                self.lbl_avisoGeracao.config(text="Processando...")
                self.mostrar_progresso()
            else:
                messagebox.showinfo("Alerta!", "CPF Inválido!") 
        
# Criar janela do Tkinter
if __name__ == "__main__":
    root = tk.Tk()
    app = genAnuencias(root)
    root.mainloop()
